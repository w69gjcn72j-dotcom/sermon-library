#!/usr/bin/env python3
"""
add-study.py — publish St Paul's Kogarah group Bible studies into the
Word & Prayer site (this repo), alongside the sermons, devotions and Seven Stars.

Usage:
    python3 add-study.py --rebuild     # scan the Sermon folder, regenerate every page + index
    python3 add-study.py --check       # verify entries <-> pages, report gaps

Collections and their page prefixes:
    kogyouth_*            -> ky-   KOGYouth Growth Groups
    homegroup_adults_*    -> gg-   Home Groups (adults)
    homegroup_youngadults_* -> ya- Young Adults (19-30)
    *Nursing Home*        -> nh-   Nursing Home Talks

Pages are flat at the repo root (ky-001-slug.html), exactly like the sermon pages.
The browse page bible-studies.html owns the STUDIES array between the
/* STUDIES:START */ and /* STUDIES:END */ markers; this script rewrites only that.
"""
import argparse, html, json, os, re, sys, unicodedata
from datetime import date

try:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))
SERMON_DIR = os.path.abspath(os.path.join(HERE, os.pardir))
INDEX_EN = os.path.join(HERE, "bible-studies.html")
INDEX_CN = os.path.join(HERE, "bible-studies-cn.html")
CN_DIR = os.path.join(HERE, "cn")      # one <pagebase>.json per translated study

COLLECTIONS = [
    ("kogyouth_",             "ky", "KOGYouth",      "KOGYouth Growth Groups",   "青少年成长小组查经"),
    ("homegroup_adults_",     "gg", "Home Groups",   "Home Groups (adults)",     "成人家庭小组查经"),
    ("homegroup_youngadults_","ya", "Young Adults",  "Young Adults (19–30)","青年成人查经"),
    ("nursing home",          "nh", "Nursing Home",  "Nursing Home Talks",       "安老院短讲"),
]
MONTHS = ["January","February","March","April","May","June","July","August",
          "September","October","November","December"]

# ---------------------------------------------------------------- helpers

def slugify(s):
    s = unicodedata.normalize("NFKD", s)
    s = s.replace("’", "").replace("'", "")
    s = re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-").lower()
    return re.sub(r"-{2,}", "-", s)[:70]


def classify(path):
    name = os.path.basename(path).lower()
    for token, prefix, short, long_, cn in COLLECTIONS:
        if token in name:
            return dict(prefix=prefix, short=short, long=long_, cn=cn)
    return None


def find_docs():
    """Every study .docx in the Sermon folder, whatever depth it sits at."""
    out = []
    for root, dirs, files in os.walk(SERMON_DIR):
        dirs[:] = [d for d in dirs if d not in (".git", "public_site", "Media & Recordings",
                                                "Recordings", "Sermons on MP3")]
        for f in files:
            if not f.endswith(".docx") or f.startswith("~$"):
                continue
            if "_handout" in f.lower():          # student print copy, not a web page
                continue
            p = os.path.join(root, f)
            if classify(p):
                out.append(p)
    return sorted(out)


def doc_date(path, dateline):
    m = re.search(r"(\d{8})", os.path.basename(path))
    if m:
        y, mo, d = int(m.group(1)[:4]), int(m.group(1)[4:6]), int(m.group(1)[6:])
        return date(y, mo, d)
    m = re.search(r"(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})", dateline or "")
    if m and m.group(2) in MONTHS:
        return date(int(m.group(3)), MONTHS.index(m.group(2)) + 1, int(m.group(1)))
    return date(1970, 1, 1)

# ---------------------------------------------------------------- parsing

def _size(par):
    for r in par.runs:
        if r.font.size:
            return r.font.size.pt
    return None


def _is_list(par):
    pPr = par._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    return "List Paragraph" in _style(par)


def _body_size(doc):
    """Most common size among ordinary (non-bold) paragraphs — heading sizes vary by doc."""
    from collections import Counter
    c = Counter()
    for par in doc.paragraphs:
        if not par.text.strip():
            continue
        if any(r.bold for r in par.runs if r.bold):
            continue
        sz = _size(par)
        if sz:
            c[sz] += 1
    return c.most_common(1)[0][0] if c else 11.0


def _style(par):
    try:
        return par.style.name or ""
    except Exception:
        return ""


NH_NOISE = re.compile(
    r"^(nursing home sermon|a shortened talk|adapted from|nursing home service|"
    r"date preached|preached\s*:|approx|rev david yung|st paul'?s anglican)", re.I)
REF_RE = re.compile(r"^[1-3]?\s*[A-Z][A-Za-z]+\s+\d+[:.]\d+([–\-]\d+)?\s*$")
BIGIDEA_RE = re.compile(r"big\s*idea\s*[::]?\s*(.+)", re.I | re.S)
DATE_RE = re.compile(r"(\d{1,2})\s+(January|February|March|April|May|June|July|August|"
                     r"September|October|November|December)\s+(\d{4})")


HEADLINE_RE = re.compile(
    r"^\s*((?:[1-3]\s*)?[A-Z][A-Za-z]+\.?\s*\d+[:.]\d+(?:\s*[–\-]\s*\d+(?::\d+)?)?)"
    r"\s*[—|·:\-–]+\s*(.+)$")


def _clean_title(t):
    t = t.strip().strip("“”\"'").strip()
    return re.sub(r"\s{2,}", " ", t)


def _split_headline(text):
    m = HEADLINE_RE.match(text)
    if m:
        return m.group(1).strip(), _clean_title(m.group(2))
    for sep in ("—", "|", "·", " – ", " - "):
        if sep in text:
            a, b = text.split(sep, 1)
            return a.strip(), _clean_title(b)
    return text.strip(), _clean_title(text)


def parse(path):
    d = docx.Document(path)
    coll = classify(path)
    base = _body_size(d)
    head_min = base + 2.5
    blocks = []          # (kind, value, size)

    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            par = Paragraph(child, d)
            text = " ".join(par.text.split())
            if not text:
                continue
            sz = _size(par)
            bold = any(r.bold for r in par.runs if r.bold)
            if sz and bold and sz >= head_min:
                blocks.append(("h", text, sz))
            elif _is_list(par):
                blocks.append(("li", text, sz))
            elif bold and len(text) < 46 and not text.endswith((".", "?", "!")):
                blocks.append(("h3", text, sz))
            else:
                blocks.append(("p", text, sz))
        elif child.tag == qn("w:tbl"):
            tb = Table(child, d)
            if len(tb.columns) == 1:
                cell = "\n".join(" ".join(p.text.split())
                                 for p in tb.rows[0].cells[0].paragraphs if p.text.strip())
                blocks.append(("box", cell, None))
            else:
                qs = []
                for row in tb.rows:
                    q = " ".join(row.cells[-1].text.split())
                    q = re.sub(r"(_{3,}\s*)+$", "", q).strip()
                    if q:
                        qs.append(q)
                if qs:
                    blocks.append(("ol", qs, None))

    title = ref = dateline = idea = ""

    if coll["prefix"] == "nh":
        # Talks have no standard header block: read everything before the Big Idea line.
        cut = len(blocks)
        for i, (k, v, s) in enumerate(blocks):
            if k == "p" and BIGIDEA_RE.match(v):
                cut = i
                break
        lead = [(v, s or 0) for k, v, s in blocks[:cut] if k in ("p", "h", "h3")]
        if cut < len(blocks):
            idea = " ".join(BIGIDEA_RE.match(blocks[cut][1]).group(1).split())
            blocks = [("box", "Big Idea: " + idea, None)] + blocks[cut + 1:]
        else:
            blocks = blocks[cut:]
        for v, s in lead:
            if not ref and REF_RE.match(v):
                ref = v
        cands = [(v, s) for v, s in lead if not NH_NOISE.match(v) and not REF_RE.match(v)]
        if cands:
            best = max(range(len(cands)), key=lambda i: cands[i][1])
            title = cands[best][0]
            if title.endswith(":") and best + 1 < len(cands):
                title = (title.rstrip(":") + ": " + cands[best + 1][0]).strip()
        for v, s in lead:
            m = DATE_RE.search(v)
            if m:
                dateline = m.group(0)
                break
        if not title:
            fn = os.path.splitext(os.path.basename(path))[0]
            fn = re.sub(r"^\[[^\]]*\]\s*", "", fn)
            fn = re.sub(r"(?i)nursing home sermon\s*[-–—:]?\s*", "", fn)
            title = _clean_title(fn.replace("- ", ": "))
        if not ref:
            ref = coll["long"]
    else:
        meta = [v for k, v, s in blocks[:3] if k in ("p", "h", "h3")][:3]
        blocks = blocks[len(meta):]
        headline = meta[1] if len(meta) > 1 else os.path.basename(path)
        dateline = meta[2] if len(meta) > 2 else ""
        ref, title = _split_headline(headline)

    if not title:
        title = os.path.splitext(os.path.basename(path))[0]
    title = _clean_title(title)

    # Big Idea: prefer the info box, fall back to any paragraph that carries one.
    if not idea:
        for k, v, s in blocks:
            if k in ("box", "p"):
                m = BIGIDEA_RE.search(v)
                if m:
                    idea = re.split(r"[\n\s]*(?:Aim|You’ll need|You'll need)\s*[::]?[\n\s]",
                                    m.group(1))[0]
                    idea = " ".join(idea.split())
                    break

    blocks = [(k, v) for k, v, s in blocks]
    when = doc_date(path, dateline)
    if not dateline:
        dateline = "%s %d" % (MONTHS[when.month - 1], when.year)
    return dict(src=path, coll=coll, title=title.strip(), ref=ref.strip(), idea=idea,
                date=when, note="%s %d" % (MONTHS[when.month - 1], when.year),
                dateline=dateline, blocks=blocks)

# ---------------------------------------------------------------- rendering

PAGE_CSS = """
  *,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
  :root{--night:#0e1a2b;--night-soft:#142339;--night-card:#1a2c47;--brass:#d4a253;
        --brass-soft:#b08539;--cream:#f5ead4;--cream-dim:#c9bfa8;
        --line:rgba(212,162,83,0.18);--line-strong:rgba(212,162,83,0.42)}
  html{font-size:18px;-webkit-text-size-adjust:100%}
  body{font-family:'Inter Tight',system-ui,-apple-system,sans-serif;background:var(--night);
       color:var(--cream);line-height:1.7;min-height:100vh;display:flex;flex-direction:column}
  .shell{width:min(720px,100%);margin:0 auto;padding:0 1.35rem;flex:1}
  .topbar{display:flex;align-items:center;justify-content:space-between;gap:1rem;
          padding:1.4rem 0 0.4rem;font-size:0.82rem}
  .topbar a{color:var(--brass);text-decoration:none;letter-spacing:0.1em;text-transform:uppercase}
  .topbar a:hover{text-decoration:underline}
  header{padding:1.8rem 0 1.6rem;border-bottom:1px solid var(--line)}
  .eyebrow{font-size:0.72rem;font-weight:500;letter-spacing:0.26em;text-transform:uppercase;
           color:var(--brass);margin-bottom:0.9rem}
  h1{font-family:'Fraunces',Georgia,serif;font-weight:600;font-size:clamp(1.8rem,5.5vw,2.5rem);
     line-height:1.15;letter-spacing:-0.01em}
  .ref{margin-top:0.7rem;color:var(--cream-dim);font-size:0.97rem}
  main{padding:1.9rem 0 1rem}
  main h2{font-family:'Fraunces',Georgia,serif;font-weight:600;font-size:1.3rem;
          color:var(--brass);margin:2.1rem 0 0.7rem;line-height:1.3}
  main h2:first-child{margin-top:0}
  main h3{font-family:'Inter Tight',system-ui,sans-serif;font-weight:600;font-size:1rem;
          color:var(--cream);margin:1.3rem 0 0.35rem;letter-spacing:0.01em}
  main p{margin:0.75rem 0;color:var(--cream)}
  main ul,main ol{margin:0.75rem 0 0.75rem 1.3rem;color:var(--cream)}
  main li{margin:0.45rem 0;padding-left:0.2rem}
  main ol{counter-reset:q;list-style:none;margin-left:0}
  main ol li{counter-increment:q;position:relative;padding:0.7rem 0.9rem 0.7rem 3rem;
             background:var(--night-soft);border:1px solid var(--line);border-radius:12px;
             margin:0.6rem 0}
  main ol li::before{content:counter(q);position:absolute;left:0.85rem;top:0.62rem;
             width:1.6rem;height:1.6rem;border-radius:50%;background:rgba(212,162,83,0.14);
             border:1px solid var(--line-strong);color:var(--brass);font-size:0.82rem;
             font-weight:600;display:grid;place-items:center}
  .box{background:rgba(20,35,57,0.7);border:1px solid var(--line);border-left:3px solid var(--brass);
       border-radius:12px;padding:1.05rem 1.2rem;margin:1.2rem 0}
  .box p{margin:0.35rem 0;font-size:0.95rem}
  .box p:first-child{margin-top:0} .box p:last-child{margin-bottom:0}
  .box b{color:var(--brass)}
  footer{border-top:1px solid var(--line);text-align:center;color:var(--cream-dim);
         font-size:0.84rem;line-height:1.85;padding:2.2rem 1.35rem 2.6rem;margin-top:2rem}
  footer a{color:var(--brass);text-decoration:none}
  footer .sdg{display:block;margin-top:0.8rem;font-family:'Fraunces',Georgia,serif;
              font-style:italic;color:var(--brass-soft)}
  @media (max-width:420px){html{font-size:17px}}
"""

PAGE_TMPL = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover">
<title>{title} — St Paul's Anglican Kogarah</title>
<meta name="description" content="{desc}">
<meta name="theme-color" content="#0e1a2b">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,400;0,9..144,600;1,9..144,400&family=Inter+Tight:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>{css}</style>
</head>
<body>
<div class="shell">
  <nav class="topbar">
    <a href="bible-studies.html">← All studies</a>
    <a href="{toggle}">{toggle_label}</a>
  </nav>
  <header>
    <p class="eyebrow">{coll}</p>
    <h1>{title}</h1>
    <p class="ref">{ref} &middot; {dateline}</p>
  </header>
  <main>
{body}
  </main>
</div>
<footer>
  Rev David Yung &middot; St Paul's Anglican Church Kogarah<br>
  <a href="mailto:dyung@kogarah.church">dyung@kogarah.church</a>
  <span class="sdg">Soli Deo Gloria</span>
</footer>
</body>
</html>
"""


def render_body(blocks):
    out, i = [], 0
    while i < len(blocks):
        kind, val = blocks[i]
        if kind == "h":
            out.append("    <h2>%s</h2>" % html.escape(val))
        elif kind == "h3":
            out.append("    <h3>%s</h3>" % html.escape(val))
        elif kind == "p":
            out.append("    <p>%s</p>" % html.escape(val))
        elif kind == "box":
            lines = []
            for ln in val.split("\n"):
                m = re.match(r"([A-Za-z’' ]{3,18}):\s*(.*)", ln)
                if m:
                    lines.append("<p><b>%s:</b> %s</p>" % (html.escape(m.group(1)), html.escape(m.group(2))))
                else:
                    lines.append("<p>%s</p>" % html.escape(ln))
            out.append('    <div class="box">%s</div>' % "".join(lines))
        elif kind == "ol":
            out.append("    <ol>")
            out += ["      <li>%s</li>" % html.escape(q) for q in val]
            out.append("    </ol>")
        elif kind == "li":
            group = []
            while i < len(blocks) and blocks[i][0] == "li":
                group.append(blocks[i][1]); i += 1
            out.append("    <ul>")
            out += ["      <li>%s</li>" % html.escape(x) for x in group]
            out.append("    </ul>")
            continue
        i += 1
    return "\n".join(out)


def render_page(e):
    desc = (e["idea"] or e["title"])[:180]
    cn = cn_path(e)
    return PAGE_TMPL.format(title=html.escape(e["title"]), desc=html.escape(desc),
                            toggle=(cn_file(e) if cn else "index.html"),
                            toggle_label=("中文" if cn else "Word &amp; Prayer"),
                            css=PAGE_CSS, coll=html.escape(e["coll"]["long"]),
                            ref=html.escape(e["ref"]),
                            dateline=html.escape(e["dateline"] or e["note"]),
                            body=render_body(e["blocks"]))

CN_PAGE_TMPL = """<!DOCTYPE html>
<html lang="zh-Hans">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover">
<title>{title} — 高嘉华圣保罗圣公会</title>
<meta name="description" content="{desc}">
<meta name="theme-color" content="#0e1a2b">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,600&family=Inter+Tight:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>{css}
  body{{font-family:'Inter Tight','PingFang SC','Songti SC','Microsoft YaHei',system-ui,sans-serif;
        line-height:1.9}}
  h1,main h2,a.study h2{{font-family:'Fraunces',Georgia,'Songti SC',serif}}
</style>
</head>
<body>
<div class="shell">
  <nav class="topbar">
    <a href="bible-studies-cn.html">← 全部查经</a>
    <a href="{en}">English</a>
  </nav>
  <header>
    <p class="eyebrow">{coll}</p>
    <h1>{title}</h1>
    <p class="ref">{ref} &middot; {dateline}</p>
  </header>
  <main>
{body}
  </main>
  <p style="margin:2rem 0 0;font-size:0.82rem;color:#8a8370;line-height:1.7">
    本页中文译自英文原稿，供弟兄姊妹参考；如与英文版有出入，以英文版为准。
  </p>
</div>
<footer>
  翁沛偉牧师 Rev David Yung &middot; 高嘉华圣保罗圣公会<br>
  <a href="mailto:dyung@kogarah.church">dyung@kogarah.church</a>
  <span class="sdg">Soli Deo Gloria</span>
</footer>
</body>
</html>
"""


def cn_file(e):
    return e["file"].replace(".html", "-cn.html")


def cn_path(e):
    f = os.path.join(CN_DIR, e["file"].split("-")[0] + "-" + e["file"].split("-")[1] + ".json")
    return f if os.path.exists(f) else None


def render_cn_page(e, data):
    blocks = [(b[0], b[1]) for b in data["blocks"]]
    desc = (data.get("idea") or data["title"])[:180]
    return CN_PAGE_TMPL.format(title=html.escape(data["title"]), desc=html.escape(desc),
                               css=PAGE_CSS, coll=html.escape(data.get("coll", e["coll"]["cn"])),
                               ref=html.escape(data.get("ref", e["ref"])),
                               dateline=html.escape(data.get("dateline", e["dateline"])),
                               en=e["file"], body=render_body(blocks))


# ---------------------------------------------------------------- index

def write_index(entries, path, marker_lang="en"):
    if not os.path.exists(path):
        sys.exit("missing browse page: %s" % path)
    src = open(path, encoding="utf-8").read()
    rows = []
    for e in entries:
        rows.append("  " + json.dumps({
            "n": e["n"], "c": e["coll"]["prefix"], "t": e["title"], "r": e["ref"],
            "d": e["date"].isoformat(), "note": e["note"], "idea": e["idea"],
            "u": e["file"],
        }, ensure_ascii=False))
    block = "/* STUDIES:START */\nconst STUDIES = [\n" + ",\n".join(rows) + "\n];\n/* STUDIES:END */"
    new = re.sub(r"/\* STUDIES:START \*/.*?/\* STUDIES:END \*/", lambda m: block, src, flags=re.S)
    if new == src and "STUDIES:START" not in src:
        sys.exit("markers not found in %s" % path)
    open(path, "w", encoding="utf-8").write(new)

# ---------------------------------------------------------------- collection pages

# Each collection gets a real page at its own URL (not a redirect) so that
# "Add to Home Screen" picks up that collection's icon rather than the
# shared one. A <base> element lets the copied markup keep working from a
# nested folder without rewriting every relative link.
COLLECTION_PAGES = [
    # (folder,                     code, source page,             icon,  manifest,                    title,                  app title)
    ("studies/home-groups",        "gg", "bible-studies.html",    "gg", "studies-home-groups.webmanifest",    "Home Group Studies",   "Home Groups"),
    ("studies/youth",              "ky", "bible-studies.html",    "ky", "studies-youth.webmanifest",          "KOGYouth Studies",     "KOGYouth"),
    ("studies/young-adults",       "ya", "bible-studies.html",    "ya", "studies-young-adults.webmanifest",   "Young Adults Studies", "Young Adults"),
    ("studies/nursing-home",       "nh", "bible-studies.html",    "nh", "studies-nursing-home.webmanifest",   "Nursing Home Talks",   "Nursing Home"),
    ("studies",                    "all","bible-studies.html",    "bs", "bible-studies.webmanifest",          "Bible Studies",        "Studies"),
    ("studies/chinese",            "all","bible-studies-cn.html", "bs", "bible-studies.webmanifest",          "查经材料",              "查经"),
    ("studies/chinese/home-groups","gg","bible-studies-cn.html",  "gg", "studies-home-groups.webmanifest",    "成人家庭小组查经",       "家庭小组"),
]


def write_collection_pages():
    made = []
    for folder, code, source, icon, manifest, title, app in COLLECTION_PAGES:
        src = open(os.path.join(HERE, source), encoding="utf-8").read()
        depth = len(folder.strip("/").split("/"))
        base = "../" * depth

        # resolve every relative URL against the site root
        out = src.replace("<head>", '<head>\n<base href="%s">' % base, 1)

        # swap in this collection's icons, manifest and titles
        out = re.sub(r'<link rel="apple-touch-icon"[^>]*>', 
                     '<link rel="apple-touch-icon" sizes="180x180" href="%s-apple-touch-icon.png">' % icon, out, count=1)
        out = re.sub(r'<link rel="icon" type="image/png" sizes="32x32"[^>]*>',
                     '<link rel="icon" type="image/png" sizes="32x32" href="%s-favicon-32.png">' % icon, out, count=1)
        out = re.sub(r'<link rel="icon" type="image/png" sizes="512x512"[^>]*>',
                     '<link rel="icon" type="image/png" sizes="512x512" href="%s-icon-512.png">' % icon, out, count=1)
        out = re.sub(r'<link rel="manifest"[^>]*>', '<link rel="manifest" href="%s">' % manifest, out, count=1)
        out = re.sub(r'<meta name="apple-mobile-web-app-title"[^>]*>',
                     '<meta name="apple-mobile-web-app-title" content="%s">' % html.escape(app), out, count=1)
        out = re.sub(r"<title>.*?</title>",
                     "<title>%s &mdash; St Paul's Anglican Kogarah</title>" % html.escape(title), out, count=1, flags=re.S)

        # open on this collection
        out = out.replace('let active = VALID.indexOf(wanted) !== -1 ? wanted : "all";',
                          'let active = VALID.indexOf(wanted) !== -1 ? wanted : "%s";' % code, 1)

        os.makedirs(os.path.join(HERE, folder), exist_ok=True)
        open(os.path.join(HERE, folder, "index.html"), "w", encoding="utf-8").write(out)
        made.append(folder + "/")
    return made


# ---------------------------------------------------------------- commands

def build():
    docs = find_docs()
    entries = [parse(p) for p in docs]
    seen = {}
    for e in entries:
        k = (e["coll"]["prefix"], e["ref"].lower(), e["title"].lower())
        cur = seen.get(k)
        if cur is None or e["date"] > cur["date"]:
            seen[k] = e
    entries = list(seen.values())
    entries.sort(key=lambda e: (e["coll"]["prefix"], e["date"], e["title"]))
    counters = {}
    for e in entries:
        pfx = e["coll"]["prefix"]
        counters[pfx] = counters.get(pfx, 0) + 1
        e["n"] = counters[pfx]
        e["file"] = "%s-%03d-%s.html" % (pfx, e["n"], slugify(e["title"]))
    return entries


def cmd_rebuild():
    entries = build()
    for e in entries:
        open(os.path.join(HERE, e["file"]), "w", encoding="utf-8").write(render_page(e))
    cn_entries = []
    for e in entries:
        src = cn_path(e)
        if not src:
            continue
        data = json.load(open(src, encoding="utf-8"))
        open(os.path.join(HERE, cn_file(e)), "w", encoding="utf-8").write(render_cn_page(e, data))
        ce = dict(e)
        ce["title"] = data["title"]
        ce["ref"] = data.get("ref", e["ref"])
        ce["idea"] = data.get("idea", "")
        ce["note"] = data.get("note", e["note"])
        ce["file"] = cn_file(e)
        cn_entries.append(ce)

    ordered = sorted(entries, key=lambda e: (e["date"], e["coll"]["prefix"]), reverse=True)
    write_index(ordered, INDEX_EN)
    write_index(sorted(cn_entries, key=lambda e: (e["date"], e["coll"]["prefix"]), reverse=True),
                INDEX_CN)
    by = {}
    for e in entries:
        by[e["coll"]["long"]] = by.get(e["coll"]["long"], 0) + 1
    print("Wrote %d study pages" % len(entries))
    for k in sorted(by):
        print("  %-26s %d" % (k, by[k]))
    print("Chinese pages: %d of %d" % (len(cn_entries), len(entries)))
    print("Collection pages: %s" % ", ".join(write_collection_pages()))
    print("Index updated: bible-studies.html, bible-studies-cn.html")


def cmd_check():
    entries = build()
    listed = set(e["file"] for e in entries)
    listed |= set(cn_file(e) for e in entries if cn_path(e))
    problems = 0
    for e in entries:
        if not os.path.exists(os.path.join(HERE, e["file"])):
            print("MISSING PAGE  %s  (%s)" % (e["file"], e["title"])); problems += 1
        if not e["idea"]:
            print("NO BIG IDEA   %s" % e["file"]); problems += 1
    for f in sorted(os.listdir(HERE)):
        if re.match(r"^(ky|gg|ya|nh)-\d{3}-", f) and f not in listed:
            print("ORPHAN PAGE   %s  (no source .docx)" % f); problems += 1
    src = open(INDEX_EN, encoding="utf-8").read()
    for e in entries:
        if e["file"] not in src:
            print("NOT IN INDEX  %s" % e["file"]); problems += 1
    cnsrc = open(INDEX_CN, encoding="utf-8").read()
    for e in entries:
        if cn_path(e) and cn_file(e) not in cnsrc:
            print("NOT IN CN INDEX  %s" % cn_file(e)); problems += 1
    todo = [e["file"] for e in entries if not cn_path(e)]
    if todo:
        print("NO CHINESE YET (%d): %s" % (len(todo), ", ".join(t[:6] for t in todo)))
    print("%d studies checked, %d problem(s)." % (len(entries), problems))
    return problems


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--rebuild", action="store_true", help="regenerate every page and the index")
    ap.add_argument("--check", action="store_true", help="verify pages and index agree")
    a = ap.parse_args()
    if a.rebuild:
        cmd_rebuild()
    elif a.check:
        sys.exit(1 if cmd_check() else 0)
    else:
        ap.print_help()
